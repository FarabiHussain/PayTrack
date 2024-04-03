import variables as vars
import docx
import os
import customtkinter as ctk
from glob import glob
from logic_records import *
from document_utils import *
from logic_gui import *
from docx.shared import Cm as CM
from datetime import datetime as dt
from path_manager import *
from CTkMessagebox import CTkMessagebox as popup


# write the invoice number and date on the top
def write_invoice_info(document, doc_id, billed_to):

    timestamp = str(dt.now().strftime("%d %B %Y, %H:%M"))

    if (len(billed_to.strip()) == 0):
        insert_paragraph_after(
            document.add_paragraph(),
            (f"Payment#\t{doc_id}\nDate\t\t{timestamp}\n\n"),
        )
    else:
        insert_paragraph_after(
            document.add_paragraph(),
            (f"Payment#\t{doc_id}\nDate\t\t{timestamp}\nBilled to\t{billed_to}\n"),
        )


# itemized table
def write_items_table(document, data):

    # Table data in a form of list
    headings = [
        {"heading": "SL.", "width": 1.0},
        {"heading": "DESCRIPTION", "width": 20.0},
        {"heading": "QTY", "width": 1.0},
        {"heading": "RATE", "width": 1.0},
        {"heading": "AMOUNT", "width": 1.0},
    ]

    # Creating a table object
    items_table = document.add_table(rows=1, cols=5)

    # add heading in the 1st row of the table
    row = items_table.rows[0].cells
    for idx, col in enumerate(headings):
        row[idx].text = col["heading"]

    # add data from the list to the table
    for index, entry in enumerate(data):

        # Adding a row and then adding data in it.
        row = items_table.add_row().cells

        row[0].text = str(index + 1)
        row[1].text = str(entry["desc"])
        row[2].text = str(entry["qty"])
        row[3].text = str(entry["rate"])
        row[4].text = str(int(entry["qty"]) * float(entry["rate"]))

    # set the table borders
    for cell in items_table.rows[0].cells:
        set_cell_border(cell, bottom={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "10"})
        set_cell_border(cell, top={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "15"})

    # set column widths
    for idx, col in enumerate(headings):
        for index, cell in enumerate(items_table.columns[idx].cells):
            cell.width = CM(col["width"])

            if index > 0:
                set_cell_border(cell, bottom={"sz": 6, "color": "#DDDDDD", "val": "single", "space": "15"},)

        set_cell_border(cell, bottom={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "0"})

    # set row heights
    for index, row in enumerate(items_table.rows):
        row.height = CM(1)

    make_rows_bold(items_table.rows[0])


# table containing taxes and total
def write_totals_table(document): 
    total_table = document.add_table(rows=1, cols=5)

    gst = vars.form['gst_display_amount'].cget("text")
    pst = vars.form['pst_display_amount'].cget("text")
    total = vars.form['total_display_amount'].cget("text")

    total_table_data = [
        [str("GST @5%"), "", "", "", gst],
        [str("PST @7%"), "", "", "", pst],
        [str("TOTAL"), "", "", "", total],
    ]

    for total_row in total_table_data:
        row = total_table.add_row().cells

        for idx, col in enumerate(total_row):
            row[idx].text = col

        make_rows_bold(total_table.rows[-1])


# combine the above functions to generate the final product
def generate_invoice(cwd):

    if len(vars.items) < 1:
        popup(title="", message='Add at least one item to create a document.', corner_radius=2)
        return False
    
    if check_special([vars.form['client_textvariable'].get()], is_string=True) is False:
        return False

    doc_id = "{:010}".format((read_from_record(cwd) + 1))
    client_name = (vars.form['client_input'].get()).strip().lower().replace(" ", "_").replace(", ",";").replace(",",";")
    write_to_record(cwd, doc_id, client_name)

    # instance of a word document using the template
    document = docx.Document(resource_path(cwd + "\\assets\\templates\\receipt.docx"))

    billed_to = vars.form['client_input'].get()

    write_invoice_info(document, doc_id, billed_to)
    write_items_table(document, vars.items)
    write_totals_table(document)

    timestamp = dt.now().strftime("%Y.%m.%d")
    billed_to = (vars.form['client_input'].get()).strip().lower().replace(" ", "_")
    filename = f"{cwd}\\output\\{timestamp}-{doc_id}.docx"

    if (len(billed_to.strip()) > 0):
        filename = f"{cwd}\\output\\{timestamp}-{doc_id}-{billed_to}.docx"

    # set up the output directory
    output_dir = os.getcwd() + "\\output\\"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # save the document with the current invoice ID
    document.save(filename.lower())
    os.startfile(filename.lower())


# open documents
def open_doc_by_filter():

    file_found = False
    doc_id = vars.search_window['doc_id_search'].get()
    client_name = vars.search_window['client_name_search'].get()

    if (len(doc_id) == 0 and len(client_name) == 0):
        popup(title="", message='Nothing entered into search fields', corner_radius=2)
        return False

    search_filters = []

    try:
        # remove surrounding whitespace, and ensure it is 10-digits long
        doc_id = "{:010}".format(int((vars.search_window['doc_id_search'].get()).strip()))
        search_filters.append(doc_id)
    except:
        print("doc_id not entered")

    try:
        # remove surrounding whitespace, convert to lowercase, replace spaces with underscores 
        client_name = (vars.search_window['client_name_search'].get()).strip().lower().replace(" ", "_")
        search_filters.append(client_name)
    except:
        print("client_name not entered")

    # try to find all matches and open them
    try:
        open_counter = 0
        open_limit = int(vars.search_window['qty_of_docs_to_open'].cget('text'))
        for f in sorted(glob(vars.cwd + "\\output\\*.docx"), key=os.path.getmtime, reverse=True):
            if all(keyword in f for keyword in search_filters):
                os.startfile(f)
                open_counter += 1
                file_found = True

            if (open_counter == open_limit):
                return True

    except Exception as e:
        print(e)

    # no match was found so display a popup message
    if not file_found:
        popup(title="", message='No match found', corner_radius=2)


# change the number of documents to be opened
def change_doc_count(amount):
    new_qty = int(vars.search_window['qty_of_docs_to_open'].cget('text')) + amount

    if (new_qty <= 1):
        new_qty = 1
        vars.search_window['minus_button'].configure(state='disabled', fg_color='#EEEEEE')
    elif new_qty >= 10:
        new_qty = 10
        vars.search_window['plus_button'].configure(state='disabled', fg_color='#EEEEEE')
    else:
        vars.search_window['minus_button'].configure(state='normal', fg_color="#23265e")
        vars.search_window['plus_button'].configure(state='normal', fg_color="#23265e")

    new_qty = "{:02}".format(new_qty)

    vars.search_window['qty_of_docs_to_open'].configure(text=new_qty)


# display the popup which allows users to search for existing documents
def search_documents():

    if (vars.search_window['popup'] is None or not vars.search_window['popup'].winfo_exists()): 
        vars.search_window['popup'] = ctk.CTkToplevel()

        w = 300
        h = 230
        x = (vars.screen_sizes['ws']/2) - (w/2)
        y = (vars.screen_sizes['hs']/2) - (h/2)

        ctk.CTkFrame(
            vars.search_window['popup'], corner_radius=2, border_width=1, width=260, height=70, fg_color='white'
        ).place(x=20, y=140)

        ctk.CTkLabel(vars.search_window['popup'], text="Document ID", bg_color='#E5E5E5', font=vars.font_family).place(x=20, y=5)
        vars.search_window['doc_id_search'] = ctk.CTkEntry(vars.search_window['popup'], width=260, border_width=1, corner_radius=2, placeholder_text="leading zeros are optional")
        vars.search_window['doc_id_search'].place(x=20, y=30)

        ctk.CTkLabel(vars.search_window['popup'], text="Client Name", bg_color='#E5E5E5', font=vars.font_family).place(x=20, y=65)
        vars.search_window['client_name_search'] = ctk.CTkEntry(vars.search_window['popup'], width=260, border_width=1, corner_radius=2, placeholder_text="full or partial name")
        vars.search_window['client_name_search'].place(x=20, y=90)

        ctk.CTkLabel(vars.search_window['popup'], text="Documents to open", bg_color='#E5E5E5', font=vars.font_family).place(x=40, y=145)
        vars.search_window['qty_of_docs_to_open'] = ctk.CTkLabel(vars.search_window['popup'], width=28, height=28, corner_radius=2, text="05", fg_color='#DDDDDD', font=vars.font_family)
        vars.search_window['qty_of_docs_to_open'].place(x=80, y=170)

        vars.search_window['minus_button'] = ctk.CTkButton(vars.search_window['popup'], text="-", border_width=0, corner_radius=2, fg_color="#23265e", command=lambda:change_doc_count(-1), height=28, width=30)
        vars.search_window['minus_button'].place(x=40, y=170)

        vars.search_window['plus_button'] = ctk.CTkButton(vars.search_window['popup'], text="+", border_width=0, corner_radius=2, fg_color="#23265e", command=lambda:change_doc_count(+1), height=28, width=30)
        vars.search_window['plus_button'].place(x=120, y=170)

        ctk.CTkButton(
            vars.search_window['popup'], text="", image=vars.icons['open'], border_width=0, corner_radius=2, fg_color="#23265e", command=lambda:open_doc_by_filter(), width=72, height=42
        ).place(x=188, y=156)

        ## render the popup
        vars.search_window['popup'].geometry('%dx%d+%d+%d' % (w, h, x, y))
        vars.search_window['popup'].resizable(False, False)
        vars.search_window['popup'].configure(fg_color='white')
       
        try:
            vars.search_window['popup'].after(201, lambda: vars.search_window['popup'].iconbitmap("assets\\icons\\logo.ico"))
        except Exception as e:
            pass
       
        vars.search_window['popup'].title("Search Payment")
        vars.search_window['popup'].after(202, lambda: vars.search_window['popup'].focus())

    else:    
        vars.search_window['popup'].focus()

