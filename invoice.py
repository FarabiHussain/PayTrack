from doc_utils import *
from data_utils import *
from docx.shared import Cm as CM
import variables as vars
from datetime import datetime as dt


#
def generate_invoice_info(document, invoice_id):

    timestamp = str(dt.now().strftime("%d/%m/%Y, %H:%M"))

    insert_paragraph_after(
        document.add_paragraph(),
        (f"Invoice#:\t{invoice_id}\nDate:\t\t{timestamp}\n"),
    )


# itemized table
def generate_items_table(document, data):

    # Table data in a form of list
    headings = [
        {"heading": "SL.", "width": 1.0},
        {"heading": "DESCRIPTION", "width": 20.0},
        {"heading": "QTY", "width": 1.0},
        {"heading": "RATE", "width": 1.0},
        {"heading": "AMOUNT", "width": 1.0},
    ]

    # Creating a table object
    items_table = document.add_table(rows=1, cols=5)

    # add heading in the 1st row of the table
    row = items_table.rows[0].cells
    for idx, col in enumerate(headings):
        row[idx].text = col["heading"]

    # add data from the list to the table
    for index, entry in enumerate(data):

        # Adding a row and then adding data in it.
        row = items_table.add_row().cells

        row[0].text = str(index + 1)
        row[1].text = str(entry["desc"])
        row[2].text = str(entry["qty"])
        row[3].text = str(entry["rate"])
        row[4].text = str(entry["qty"] * entry["rate"])

    # set the table borders
    for cell in items_table.rows[0].cells:
        set_cell_border(cell, bottom={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "10"})
        set_cell_border(cell, top={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "15"})

    # set column widths
    for idx, col in enumerate(headings):
        for index, cell in enumerate(items_table.columns[idx].cells):
            cell.width = CM(col["width"])

            if index > 0:
                set_cell_border(cell, bottom={"sz": 6, "color": "#DDDDDD", "val": "single", "space": "15"},)

        set_cell_border(cell, bottom={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "0"})

    # set row heights
    for index, row in enumerate(items_table.rows):
        row.height = CM(1)

    make_rows_bold(items_table.rows[0])


# table containing taxes and total
def generate_totals_table(document, totals): 
    total_table = document.add_table(rows=1, cols=5)

    total_table_data = [
        [str("GST @5%"), "", "", "", str(totals["gst"])],
        [str("PST @7%"), "", "", "", str(totals["pst"])],
        [str("TOTAL"), "", "", "", str(totals["dollar"])],
    ]

    for total_row in total_table_data:
        row = total_table.add_row().cells

        for idx, col in enumerate(total_row):
            row[idx].text = col

        make_rows_bold(total_table.rows[-1])


