import variables as vars
import docx
import os
from logic_records import *
from doc_utils import *
from docx.shared import Cm as CM
from datetime import datetime as dt
from path_manager import *


# write the invoice number and date on the top
def write_invoice_info(document, doc_id):

    timestamp = str(dt.now().strftime("%d/%m/%Y, %H:%M"))

    insert_paragraph_after(
        document.add_paragraph(),
        (f"Invoice#:\t{doc_id}\nDate:\t\t{timestamp}\n"),
    )


# itemized table
def write_items_table(document, data):

    # Table data in a form of list
    headings = [
        {"heading": "SL.", "width": 1.0},
        {"heading": "DESCRIPTION", "width": 20.0},
        {"heading": "QTY", "width": 1.0},
        {"heading": "RATE", "width": 1.0},
        {"heading": "AMOUNT", "width": 1.0},
    ]

    # Creating a table object
    items_table = document.add_table(rows=1, cols=5)

    # add heading in the 1st row of the table
    row = items_table.rows[0].cells
    for idx, col in enumerate(headings):
        row[idx].text = col["heading"]

    # add data from the list to the table
    for index, entry in enumerate(data):

        # Adding a row and then adding data in it.
        row = items_table.add_row().cells

        row[0].text = str(index + 1)
        row[1].text = str(entry["desc"])
        row[2].text = str(entry["qty"])
        row[3].text = str(entry["rate"])
        row[4].text = str(int(entry["qty"]) * float(entry["rate"]))

    # set the table borders
    for cell in items_table.rows[0].cells:
        set_cell_border(cell, bottom={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "10"})
        set_cell_border(cell, top={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "15"})

    # set column widths
    for idx, col in enumerate(headings):
        for index, cell in enumerate(items_table.columns[idx].cells):
            cell.width = CM(col["width"])

            if index > 0:
                set_cell_border(cell, bottom={"sz": 6, "color": "#DDDDDD", "val": "single", "space": "15"},)

        set_cell_border(cell, bottom={"sz": 6, "color": "#AAAAAA", "val": "single", "space": "0"})

    # set row heights
    for index, row in enumerate(items_table.rows):
        row.height = CM(1)

    make_rows_bold(items_table.rows[0])


# table containing taxes and total
def write_totals_table(document): 
    total_table = document.add_table(rows=1, cols=5)

    gst = vars.form['gst_display_amount'].cget("text")
    pst = vars.form['pst_display_amount'].cget("text")
    total = vars.form['total_display_amount'].cget("text")

    total_table_data = [
        [str("GST @5%"), "", "", "", gst],
        [str("PST @7%"), "", "", "", pst],
        [str("TOTAL"), "", "", "", total],
    ]

    for total_row in total_table_data:
        row = total_table.add_row().cells

        for idx, col in enumerate(total_row):
            row[idx].text = col

        make_rows_bold(total_table.rows[-1])


# combine the above functions to generate the final product
def generate_invoice(cwd):
    doc_id = "{:010}".format((read_from_record(cwd) + 1))
    write_to_record(cwd, doc_id)

    # instance of a word document using the template
    document = docx.Document(resource_path(cwd + "\\assets\\templates\\receipt.docx"))

    write_invoice_info(document, doc_id)
    write_items_table(document, vars.items)
    write_totals_table(document)

    # save the document with the current invoice ID
    document.save(cwd + "\\output\\receipt_" + doc_id + ".docx")
    os.startfile(cwd + "\\output\\receipt_" + doc_id + ".docx")

